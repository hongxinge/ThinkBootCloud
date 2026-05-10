#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
ThinkBootCloud 推广文档生成脚本
生成可直接导入各大平台的 Word 宣传文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

def set_run_format(run, font_name=None, font_size=None, bold=False, color=None, italic=False):
    """设置文字格式"""
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_styled_paragraph(doc, text, font_size=12, bold=False, color=None, alignment=None, font_name='微软雅黑', space_before=0, space_after=6):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_format(run, font_name=font_name, font_size=font_size, bold=bold, color=color)
    return p

def add_multi_format_paragraph(doc, segments, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    """添加包含多种格式文字的段落
    segments: list of tuples (text, font_size, bold, color, font_name)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        text = seg[0]
        font_size = seg[1] if len(seg) > 1 else 12
        bold = seg[2] if len(seg) > 2 else False
        color = seg[3] if len(seg) > 3 else None
        font_name = seg[4] if len(seg) > 4 else '微软雅黑'
        run = p.add_run(text)
        set_run_format(run, font_name=font_name, font_size=font_size, bold=bold, color=color)
    return p

def create_promotion_doc():
    """创建推广文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(12)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ==================== 封面标题 ====================
    doc.add_paragraph()  # 空行
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(10)
    run = p_title.add_run('ThinkBootCloud')
    set_run_format(run, font_name='微软雅黑', font_size=36, bold=True, color=(102, 126, 234))

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    run = p_sub.add_run('轻量级微服务开发框架')
    set_run_format(run, font_name='微软雅黑', font_size=24, bold=False, color=(118, 75, 162))

    # 宣传语
    p_slogan = doc.add_paragraph()
    p_slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_slogan.paragraph_format.space_after = Pt(5)
    run = p_slogan.add_run('专为 C 端客户端设计 · 开箱即用 · 一行配置即可开发')
    set_run_format(run, font_name='微软雅黑', font_size=14, bold=False, color=(100, 100, 100))

    # 协议信息
    p_protocol = doc.add_paragraph()
    p_protocol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_protocol.paragraph_format.space_after = Pt(20)
    run = p_protocol.add_run('MIT 开源协议 · 完全免费 · 可自由商用')
    set_run_format(run, font_name='微软雅黑', font_size=12, bold=True, color=(76, 175, 80))

    # 分隔线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run('━' * 30)
    set_run_format(run, font_size=10, color=(200, 200, 200))

    # ==================== 一、项目简介 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('一、', 18, True, (102, 126, 234)),
        ('项目简介', 18, True, (51, 51, 51))
    ], space_after=12)

    add_styled_paragraph(doc, 'ThinkBootCloud 是一款基于 Spring Cloud Alibaba + Spring Boot 3 的轻量级微服务开发框架。', 13, False, (80, 80, 80), space_after=8)

    add_styled_paragraph(doc, '与若依等后台管理系统不同，本框架专注于 C 端客户端场景（移动端App、Web前端、小程序等），去除了复杂的角色权限体系，仅保留基础的 Token 验证，让开发者能够开箱即用，快速开发业务逻辑。', 13, False, (80, 80, 80), space_after=10)

    # ==================== 二、核心优势 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('二、', 18, True, (102, 126, 234)),
        ('核心优势', 18, True, (51, 51, 51))
    ], space_after=12)

    advantages = [
        ('⚡ 开箱即用', '引入依赖，简单配置即可开始开发，无需从零搭建脚手架'),
        ('🔐 安全优先', '默认所有接口需要 Token 认证，明确标记才放行，杜绝安全漏洞'),
        ('📦 模块化设计', '按需引入模块，不臃肿，轻量灵活'),
        ('🚀 代码生成器', '基于数据库表一键生成 Entity、Mapper、Service、Controller'),
        ('🔄 约定优于配置', '提供合理默认值，减少配置工作量'),
        ('👨‍💻 开发者友好', '统一响应格式、全局异常处理、自动装配、全中文文档'),
    ]

    for title, desc in advantages:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(title + '\n')
        set_run_format(run, font_name='微软雅黑', font_size=14, bold=True, color=(102, 126, 234))
        run2 = p.add_run('    ' + desc)
        set_run_format(run2, font_name='微软雅黑', font_size=12, color=(100, 100, 100))

    # ==================== 三、技术栈 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('三、', 18, True, (102, 126, 234)),
        ('技术栈', 18, True, (51, 51, 51))
    ], space_after=12)

    # 技术栈表格
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头样式
    header_cells = table.rows[0].cells
    headers = ['组件', '版本', '说明']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_format(run, font_name='微软雅黑', font_size=12, bold=True, color=(255, 255, 255))
        # 设置表头背景
        shading = header_cells[i]._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '667EEA',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    tech_stack = [
        ('Spring Boot', '3.2.5', '核心框架'),
        ('Spring Cloud', '2023.0.1', '微服务框架'),
        ('Spring Cloud Alibaba', '2023.0.1.0', '阿里微服务组件'),
        ('Nacos', '2.x', '注册中心 + 配置中心'),
        ('MyBatis-Plus', '3.5.6', 'ORM框架'),
        ('Redis (Redisson)', '3.27.2', '缓存 + 分布式锁'),
        ('JWT (JJWT)', '0.12.5', 'Token认证'),
        ('OpenFeign', '4.1.x', '服务间通信'),
        ('Sentinel', '2023.0.1.0', '限流熔断'),
        ('Gateway', '4.1.x', 'API网关'),
        ('Druid', '1.2.21', '数据库连接池'),
        ('Knife4j', '4.4.0', 'API文档'),
    ]

    for tech, version, desc in tech_stack:
        row = table.add_row()
        row.cells[0].text = tech
        row.cells[1].text = version
        row.cells[2].text = desc
        for i, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_format(run, font_name='微软雅黑', font_size=11, color=(80, 80, 80))

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(5)

    # ==================== 四、模块结构 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('四、', 18, True, (102, 126, 234)),
        ('模块结构', 18, True, (51, 51, 51))
    ], space_after=12)

    modules = [
        ('think-boot-common', '公共基础模块', '统一响应R、分页支持、业务异常、全局异常处理'),
        ('think-boot-core', '核心配置模块', 'CORS配置、Jackson时间序列化、请求日志'),
        ('think-boot-auth', '认证模块', 'JWT Token生成/验证、@IgnoreAuth注解、拦截器'),
        ('think-boot-feign', '服务通信模块', 'OpenFeign集成、Token自动传递、全局日志'),
        ('think-boot-sentinel', '限流熔断模块', 'Sentinel集成、自定义限流响应'),
        ('think-boot-gateway', 'API网关模块', '网关路由、Token验证、CORS、全局错误处理'),
        ('think-boot-nacos', '服务注册模块', 'Nacos服务发现、动态配置'),
        ('think-boot-mybatis', '数据库模块', 'MyBatis-Plus、BaseEntity、分页、自动填充'),
        ('think-boot-redis', '缓存模块', 'RedisTemplate封装、Redisson分布式锁'),
        ('think-boot-file', '文件上传模块', '本地存储、阿里云OSS、文件上传下载'),
        ('think-boot-codegen', '代码生成器模块', '基于数据库表自动生成CRUD代码'),
        ('think-boot-example', '示例模块', '完整演示框架用法'),
    ]

    for mod, name, desc in modules:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f'📁 {mod}')
        set_run_format(run, font_name='Consolas', font_size=12, bold=True, color=(102, 126, 234))
        run2 = p.add_run(f' - {name}：{desc}')
        set_run_format(run2, font_name='微软雅黑', font_size=11, color=(80, 80, 80))

    # ==================== 五、快速开始 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('五、', 18, True, (102, 126, 234)),
        ('快速开始', 18, True, (51, 51, 51))
    ], space_after=12)

    steps = [
        ('第一步：引入父POM', '''<parent>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-cloud</artifactId>
    <version>1.0.0</version>
</parent>'''),
        ('第二步：引入依赖', '''<!-- 必选 -->
<dependency>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-common</artifactId>
</dependency>
<dependency>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-core</artifactId>
</dependency>
<dependency>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-auth</artifactId>
</dependency>

<!-- 按需引入 -->
<dependency>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-mybatis</artifactId>
</dependency>
<dependency>
    <groupId>com.thinkboot</groupId>
    <artifactId>think-boot-redis</artifactId>
</dependency>'''),
        ('第三步：配置 application.yml', '''server:
  port: 8080

spring:
  datasource:
    url: jdbc:mysql://localhost:3306/your_db
    username: root
    password: your_password

thinkboot:
  auth:
    jwt:
      secret: dGhpbmstYm9vdC1qd3Qtc2VjcmV0...
      skip-paths:
        - /api/auth/login
        - /api/auth/register'''),
    ]

    for title, code in steps:
        add_styled_paragraph(doc, title, 14, True, (51, 51, 51), space_before=12, space_after=8)
        # 代码块
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        # 使用浅灰背景模拟代码块
        run = p.add_run(code)
        set_run_format(run, font_name='Consolas', font_size=10, color=(50, 50, 50))

    # ==================== 六、代码示例 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('六、', 18, True, (102, 126, 234)),
        ('代码示例', 18, True, (51, 51, 51))
    ], space_after=12)

    add_styled_paragraph(doc, '6.1 创建 Controller（默认需要登录）', 14, True, (51, 51, 51), space_before=12, space_after=8)
    
    code1 = '''@RestController
@RequestMapping("/api/users")
public class UserController {
    private final UserService userService;
    
    // 默认需要登录（无需任何注解）
    @GetMapping("/{id}")
    public R<User> getById(@PathVariable Long id) {
        return R.success(userService.getById(id));
    }
    
    // 不需要登录（使用 @IgnoreAuth 注解）
    @IgnoreAuth
    @GetMapping("/public/info")
    public R<String> publicInfo() {
        return R.success("这是公开信息");
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run(code1)
    set_run_format(run, font_name='Consolas', font_size=10, color=(50, 50, 50))

    add_styled_paragraph(doc, '\n6.2 文件上传', 14, True, (51, 51, 51), space_before=12, space_after=8)
    
    code2 = '''@PostMapping("/upload")
public R<FileUploadResult> upload(@RequestParam("file") MultipartFile file) {
    FileUploadResult result = fileStorageService.upload(file);
    return R.success(result);
}'''
    p = doc.add_paragraph()
    run = p.add_run(code2)
    set_run_format(run, font_name='Consolas', font_size=10, color=(50, 50, 50))

    # ==================== 七、项目地址 ====================
    doc.add_paragraph()
    add_multi_format_paragraph(doc, [
        ('七、', 18, True, (102, 126, 234)),
        ('项目地址', 18, True, (51, 51, 51))
    ], space_after=12)

    add_styled_paragraph(doc, '🔗 Gitee 仓库：https://gitee.com/hongxinge/think-boot-cloud', 13, True, (102, 126, 234), WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_styled_paragraph(doc, '📮 问题反馈：https://gitee.com/hongxinge/think-boot-cloud/issues', 13, False, (100, 100, 100), WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # 底部装饰线
    doc.add_paragraph()
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run('━' * 40)
    set_run_format(run, font_size=10, color=(200, 200, 200))

    add_styled_paragraph(doc, 'MIT License · 完全免费 · 可自由商用', 11, True, (76, 175, 80), WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=5)
    add_styled_paragraph(doc, f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', 10, False, (150, 150, 150), WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # 保存文档
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'promotion', 'ThinkBootCloud推广文档.docx')
    doc.save(output_path)
    print(f'✅ Word 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_promotion_doc()
